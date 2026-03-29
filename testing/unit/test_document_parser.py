"""
Unit tests for DocumentParser service.

Tests file parsing, text extraction, text cleaning, and error handling.
"""

import pytest
import os
import tempfile
from pathlib import Path
from io import BytesIO
from unittest.mock import patch, MagicMock
from app.services.document_parser import DocumentParser


class TestDocumentParserFileSupport:
    """Tests for file format support."""

    def test_is_supported_pdf(self):
        """Test that PDF files are supported."""
        assert DocumentParser.is_supported("resume.pdf") is True

    def test_is_supported_docx(self):
        """Test that DOCX files are supported."""
        assert DocumentParser.is_supported("resume.docx") is True

    def test_is_supported_doc(self):
        """Test that old DOC format is not supported (only .pdf and .docx)."""
                                                                   
        assert DocumentParser.is_supported("resume.doc") is False

    def test_is_supported_txt(self):
        """Test that TXT files are not supported."""
        assert DocumentParser.is_supported("resume.txt") is False

    def test_is_supported_case_insensitive(self):
        """Test that extension check is case-insensitive."""
        assert DocumentParser.is_supported("resume.PDF") is True
        assert DocumentParser.is_supported("resume.DOCX") is True
        assert DocumentParser.is_supported("resume.Pdf") is True


class TestDocumentParserParseFile:
    """Tests for main parse_file method."""

    def test_parse_file_nonexistent_raises_error(self):
        """Test that parsing nonexistent file raises FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            DocumentParser.parse_file("/nonexistent/path/resume.pdf")

    def test_parse_file_unsupported_format_raises_error(self):
        """Test that unsupported format raises ValueError."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            temp_path = f.name
        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                DocumentParser.parse_file(temp_path)
        finally:
            os.unlink(temp_path)

    def test_parse_file_doc_raises_not_implemented(self):
        """Test that .doc format raises NotImplementedError."""
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            temp_path = f.name
        try:
            with pytest.raises(NotImplementedError, match="Legacy .doc format"):
                DocumentParser.parse_file(temp_path)
        finally:
            os.unlink(temp_path)


class TestDocumentParserCleanText:
    """Tests for text cleaning utility."""

    def test_clean_text_strips_whitespace(self):
        """Test that clean_text strips leading/trailing whitespace."""
        raw = "  \n  Hello world  \n  "
        cleaned = DocumentParser.clean_text(raw)
        assert cleaned == "Hello world"

    def test_clean_text_collapses_blank_lines(self):
        """Test that clean_text collapses multiple blank lines."""
        raw = "Line 1\n\n\n\nLine 2"
        cleaned = DocumentParser.clean_text(raw)
        assert "\n\n\n" not in cleaned
        assert "Line 1\n\nLine 2" in cleaned

    def test_clean_text_preserves_content_lines(self):
        """Test that clean_text preserves non-empty lines."""
        raw = "Line 1\nLine 2\nLine 3"
        cleaned = DocumentParser.clean_text(raw)
        assert "Line 1" in cleaned
        assert "Line 2" in cleaned
        assert "Line 3" in cleaned

    def test_clean_text_multiple_spaces_preserved(self):
        """Test that multiple spaces within text are mostly preserved."""
        raw = "Multiple  spaces  here"
        cleaned = DocumentParser.clean_text(raw)
                                                                                   
        assert "spac" in cleaned

    def test_clean_text_tabs_preserved(self):
        """Test that tabs are preserved."""
        raw = "Column1\tColumn2"
        cleaned = DocumentParser.clean_text(raw)
        assert "\t" in cleaned or "Column1" in cleaned

    def test_clean_text_empty_string(self):
        """Test cleaning empty string."""
        result = DocumentParser.clean_text("")
        assert result == ""

    def test_clean_text_only_whitespace(self):
        """Test cleaning whitespace-only string."""
        result = DocumentParser.clean_text("   \n\n  \t  ")
        assert result == ""

    def test_clean_text_multiline_paragraphs(self):
        """Test cleaning multiline paragraphs."""
        raw = """
        Experience Section
        Company A (2020-2021)
        - Worked on project X
        - Delivered feature Y

        Company B (2021-2023)
        - Led team of 5
        """
        cleaned = DocumentParser.clean_text(raw)
        assert "Experience Section" in cleaned
        assert "Company A" in cleaned
        assert "project X" in cleaned

    def test_clean_text_removes_control_characters(self):
        """Test that control characters are removed."""
        raw = "Hello\x00World\x01Test"
        cleaned = DocumentParser.clean_text(raw)
                                              
        assert "\x00" not in cleaned
        assert "\x01" not in cleaned

    def test_clean_text_preserves_meaningful_newlines(self):
        """Test that meaningful newlines structure is preserved."""
        raw = "First section\nSecond section\nThird section"
        cleaned = DocumentParser.clean_text(raw)
        lines = cleaned.split("\n")
        assert len(lines) == 3


class TestDocumentParserPdfParsing:
    """Tests for PDF parsing (pdfplumber/PyPDF2)."""

    def test_parse_pdf_empty_file_returns_string(self):
        """Test that parsing invalid PDF returns empty string or raises."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"not a pdf")
            temp_path = f.name
        try:
            result = DocumentParser.parse_pdf(temp_path)
                                                          
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_parse_pdf_nonexistent_returns_empty(self):
        """Test that parsing nonexistent PDF is handled gracefully."""
                                                                                    
        result = DocumentParser.parse_pdf("/nonexistent/file.pdf")
        assert isinstance(result, str)


class TestDocumentParserDocxParsing:
    """Tests for DOCX parsing."""

    def test_parse_docx_nonexistent_file(self):
        """Test parsing nonexistent DOCX file."""
                                                                
        with pytest.raises((FileNotFoundError, ValueError)):
            DocumentParser.parse_docx("/nonexistent/file.docx")

    def test_parse_docx_invalid_file(self):
        """Test parsing invalid DOCX file."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"not a docx file")
            temp_path = f.name
        try:
            with pytest.raises(ValueError, match="Could not parse DOCX"):
                DocumentParser.parse_docx(temp_path)
        finally:
            os.unlink(temp_path)

    def test_parse_docx_with_python_docx_installed(self):
        """Test DOCX parsing when python-docx is available."""
        try:
            from docx import Document
                                          
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name
            try:
                doc = Document()
                doc.add_paragraph("Test content")
                doc.save(temp_path)
                
                result = DocumentParser.parse_docx(temp_path)
                assert isinstance(result, str)
                assert "Test content" in result
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("python-docx not installed")


class TestDocumentParserDocUnsupported:
    """Tests for unsupported .doc format."""

    def test_parse_doc_raises_not_implemented(self):
        """Test that parsing .doc format raises NotImplementedError."""
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            temp_path = f.name
        try:
            with pytest.raises(NotImplementedError, match="Legacy .doc format"):
                DocumentParser.parse_file(temp_path)
        finally:
            os.unlink(temp_path)


class TestDocumentParserPdfPrivateMethods:
    """Tests specifically for private PDF parsing methods."""

    def test_parse_pdf_pdfplumber_with_valid_pdf(self):
        """Test PDF parsing with pdfplumber when available."""
        try:
            import pdfplumber
                                        
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                                            
                f.write(b"%PDF-1.0\n1 0 obj\n<< >>\nendobj\n")
                temp_path = f.name
            
            try:
                result = DocumentParser._parse_pdf_pdfplumber(temp_path)
                                                                  
                assert isinstance(result, str)
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("pdfplumber not installed")

    def test_parse_pdf_pypdf2_fallback(self):
        """Test PDF parsing fallback to PyPDF2."""
        try:
            import PyPDF2
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                                         
                f.write(b"%PDF-1.0\n")
                temp_path = f.name
            
            try:
                result = DocumentParser._parse_pdf_pypdf2(temp_path)
                assert isinstance(result, str)
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("PyPDF2 not installed")

    def test_parse_pdf_with_corrupted_pdf(self):
        """Test handling of corrupted PDF."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"This is not a PDF at all")
            temp_path = f.name
        
        try:
                                                                   
            result1 = DocumentParser._parse_pdf_pdfplumber(temp_path)
            result2 = DocumentParser._parse_pdf_pypdf2(temp_path)
            
                                                    
            assert isinstance(result1, str)
            assert isinstance(result2, str)
        finally:
            os.unlink(temp_path)


class TestDocumentParserDocxPrivateMethods:
    """Tests for DOCX private parsing methods."""

    def test_parse_docx_with_tables(self):
        """Test DOCX parsing with tables."""
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name
            
            try:
                doc = Document()
                doc.add_paragraph("Introduction")
                
                                        
                table = doc.add_table(rows=2, cols=2)
                table.cell(0, 0).text = "Header 1"
                table.cell(0, 1).text = "Header 2"
                table.cell(1, 0).text = "Data 1"
                table.cell(1, 1).text = "Data 2"
                
                doc.save(temp_path)
                
                result = DocumentParser.parse_docx(temp_path)
                
                                              
                assert "Header 1" in result or "Data 1" in result or len(result) > 0
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_docx_empty_document(self):
        """Test parsing empty DOCX document."""
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name
            
            try:
                doc = Document()
                doc.save(temp_path)
                
                result = DocumentParser.parse_docx(temp_path)
                                                                    
                assert isinstance(result, str)
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("python-docx not installed")


class TestDocumentParserExtensionHandling:
    """Tests for file extension handling."""

    def test_parse_file_maintains_extension_case(self):
        """Test that parse_file handles various case combinations."""
        with tempfile.NamedTemporaryFile(suffix=".PDF", delete=False) as f:
            f.write(b"%PDF-1.0\n")
            temp_path = f.name
        
        try:
            result = DocumentParser.parse_file(temp_path)
                                               
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_parse_file_mixed_case_extension(self):
        """Test file with mixed case extension."""
        with tempfile.NamedTemporaryFile(suffix=".Pdf", delete=False) as f:
            f.write(b"%PDF-1.0\n")
            temp_path = f.name
        
        try:
            result = DocumentParser.parse_file(temp_path)
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)


class TestDocumentParserPdfException:
    """Tests for PDF parsing exception handling."""

    def test_parse_pdf_when_pdfplumber_fails(self):
        """Test PDF parsing when pdfplumber raises exception."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                                        
            f.write(b"PDF%\x00\xff\xfe")                                 
            temp_path = f.name
        
        try:
                                                                      
            result = DocumentParser.parse_pdf(temp_path)
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_parse_pdf_pdfplumber_with_exception(self):
        """Test _parse_pdf_pdfplumber exception handling."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                                                              
            f.write(b"\x00\x01\x02\x03\x04\x05\x06\x07")
            temp_path = f.name
        
        try:
            result = DocumentParser._parse_pdf_pdfplumber(temp_path)
                                                   
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_parse_pdf_pypdf2_with_exception(self):
        """Test _parse_pdf_pypdf2 exception handling."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                                                          
            f.write(b"\x00\x01\x02\x03\x04\x05\x06\x07")
            temp_path = f.name
        
        try:
            result = DocumentParser._parse_pdf_pypdf2(temp_path)
                                                   
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)


class TestDocumentParserDocxException:
    """Tests for DOCX parsing exception handling."""

    def test_parse_docx_with_invalid_file(self):
        """Test DOCX parsing with invalid/corrupted DOCX file."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                                
            f.write(b"This is not a valid DOCX file")
            temp_path = f.name
        
        try:
                                                        
            with pytest.raises(ValueError):
                DocumentParser.parse_docx(temp_path)
        finally:
            os.unlink(temp_path)

    def test_parse_docx_with_empty_text(self):
        """Test DOCX parsing with all empty paragraphs and tables."""
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name
            
            try:
                doc = Document()
                                      
                doc.add_paragraph("")
                doc.add_paragraph("  ")
                doc.save(temp_path)
                
                result = DocumentParser.parse_docx(temp_path)
                                                               
                assert result == ""
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("python-docx not installed")
    """Integration tests for DocumentParser."""

    def test_parse_file_dispatch_to_pdf(self):
        """Test that PDF files are dispatched to PDF parser."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"fake pdf")
            temp_path = f.name
        try:
            result = DocumentParser.parse_file(temp_path)
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_parse_file_dispatch_to_docx(self):
        """Test that DOCX files are dispatched to DOCX parser."""
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name
            try:
                doc = Document()
                doc.add_paragraph("Sample resume content")
                doc.save(temp_path)
                
                result = DocumentParser.parse_file(temp_path)
                assert isinstance(result, str)
                assert "Sample resume content" in result
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_clean_text_integration_with_parsed_content(self):
        """Test cleaning parsed resume content."""
        raw_resume = """
        
        John Doe
        Software Engineer
        
        
        Experience
        Company A (2020-2021)
        - Python backend development
        - REST API design
        
        
        Education
        BS Computer Science (2019)
        """
        
        cleaned = DocumentParser.clean_text(raw_resume)
        
                                                               
        assert "John Doe" in cleaned
        assert "Company A" in cleaned
        assert "Education" in cleaned
        assert "BS Computer Science" in cleaned
                                              
        assert "\n\n\n" not in cleaned


class TestDocumentParserEdgeCases:
    """Edge case tests for DocumentParser."""

    def test_clean_text_with_special_characters(self):
        """Test cleaning text with special characters."""
        raw = "Name: José\nEmail: test@example.com\nSkill: C++"
        cleaned = DocumentParser.clean_text(raw)
        assert "José" in cleaned or "Jose" in cleaned
        assert "@" in cleaned

    def test_parse_file_with_valid_pdf(self):
        """Test parsing a valid PDF file."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                               
            f.write(b"%PDF-1.0\n")
            temp_path = f.name
        
        try:
            result = DocumentParser.parse_file(temp_path)
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_parse_file_nonexistent(self):
        """Test parsing nonexistent file."""
        with pytest.raises(FileNotFoundError):
            DocumentParser.parse_file("/nonexistent/path/resume.pdf")

    def test_parse_file_unsupported_format(self):
        """Test parsing unsupported file type."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"Some text content")
            temp_path = f.name
        
        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                DocumentParser.parse_file(temp_path)
        finally:
            os.unlink(temp_path)

    def test_parse_doc_unsupported(self):
        """Test parsing legacy .doc format."""
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(b"Legacy format")
            temp_path = f.name
        
        try:
            with pytest.raises(NotImplementedError):
                DocumentParser.parse_file(temp_path)
        finally:
            os.unlink(temp_path)

    def test_clean_text_with_control_characters(self):
        """Test that clean_text removes control characters."""
        raw = "Name\x00John\x01Doe\x02Resume"
        cleaned = DocumentParser.clean_text(raw)
                                              
        assert "\x00" not in cleaned
        assert "\x01" not in cleaned
        assert "\x02" not in cleaned

    def test_clean_text_preserves_tabs(self):
        """Test that clean_text preserves tabs."""
        raw = "Item1\tValue1\nItem2\tValue2"
        cleaned = DocumentParser.clean_text(raw)
                                  
        assert "\t" in cleaned

    def test_parse_docx_with_multiple_paragraphs(self):
        """Test DOCX parsing with multiple paragraphs."""
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name
            
            try:
                doc = Document()
                doc.add_paragraph("Paragraph 1")
                doc.add_paragraph("Paragraph 2")
                doc.add_paragraph("Paragraph 3")
                doc.save(temp_path)
                
                result = DocumentParser.parse_docx(temp_path)
                                                  
                assert "Paragraph 1" in result
                assert "Paragraph 2" in result
                assert "Paragraph 3" in result
            finally:
                os.unlink(temp_path)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_clean_text_with_unicode(self):
        """Test cleaning text with unicode characters."""
        raw = "Fluent in: Español, Português, 中文"
        cleaned = DocumentParser.clean_text(raw)
        assert "Fluent in:" in cleaned

    def test_parse_empty_pdf_file(self):
        """Test parsing empty PDF file."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                                               
            f.write(b"%PDF-1.0")
            temp_path = f.name
        try:
            result = DocumentParser.parse_pdf(temp_path)
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_is_supported_with_path_object(self):
        """Test is_supported works with Path-like objects."""
        from pathlib import Path
        result = DocumentParser.is_supported("resume.pdf")
        assert result is True

    def test_parse_file_with_absolute_path(self):
        """Test parse_file with absolute path."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"%PDF-1.0")
            temp_path = f.name
        try:
            abs_path = os.path.abspath(temp_path)
            result = DocumentParser.parse_file(abs_path)
            assert isinstance(result, str)
        finally:
            os.unlink(temp_path)

    def test_clean_text_preserves_list_format(self):
        """Test that clean_text preserves bullet point formatting."""
        raw = """
        Skills:
        - Python
        - JavaScript
        - Docker
        """
        cleaned = DocumentParser.clean_text(raw)
        assert "Skills:" in cleaned
        assert "Python" in cleaned
        assert "JavaScript" in cleaned
        assert "Docker" in cleaned

    def test_clean_text_multiple_sections(self):
        """Test cleaning resume with multiple distinct sections."""
        raw = """
        RESUME
        
        OBJECTIVE
        To obtain a position...
        
        SKILLS
        Python, Java, SQL
        
        EXPERIENCE
        Company A (2020-2022)
        """
        cleaned = DocumentParser.clean_text(raw)
        assert "RESUME" in cleaned
        assert "OBJECTIVE" in cleaned
        assert "SKILLS" in cleaned
        assert "EXPERIENCE" in cleaned
                                                                      
        assert "\n\n\n" not in cleaned
