"""
app/services/document_parser.py

Document Parser - Text extraction from PDF and DOCX resume files.

Provides a single entry-point (DocumentParser.parse_file) that auto-detects
the file format and delegates to the appropriate back-end.

Supported formats
─────────────────
  .pdf   pdfplumber  (primary)  →  PyPDF2  (fallback)
  .docx  python-docx  (paragraphs + table cells)
  .doc   raises NotImplementedError with user-friendly guidance

Import pattern (from any service):
    from app.services.document_parser import DocumentParser

    text = DocumentParser.parse_file("/path/to/resume.pdf")
"""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Extract plain text from uploaded resume files.

    All public methods return a stripped str.  An empty string is returned
    (not raised) when the document contains no extractable text, so callers
    can degrade gracefully without try/except at every call site.
    """

                                                                        
                
                                                                        

    @classmethod
    def parse_file(cls, file_path: str) -> str:
        """
        Auto-detect file type and return extracted plain text.

        :param file_path: Absolute or relative path to the resume file.
        :type file_path: str
        :returns: Extracted text as a single string (may be empty).
        :rtype: str
        :raises FileNotFoundError: File does not exist on disk.
        :raises ValueError: Unsupported file format.
        :raises NotImplementedError: Legacy .doc format requested.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        ext = path.suffix.lower()
        dispatch = {
            ".pdf":  cls.parse_pdf,
            ".docx": cls.parse_docx,
            ".doc":  cls._parse_doc_unsupported,
        }

        handler = dispatch.get(ext)
        if handler is None:
            raise ValueError(
                f"Unsupported file format '{ext}'. Accepted: .pdf, .docx"
            )
        return handler(file_path)

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from a PDF.

        Tries pdfplumber first (better layout handling), then falls back to
        PyPDF2.  Returns an empty string if both fail – never raises.

        :param file_path: Path to the PDF file.
        :type file_path: str
        :returns: Extracted text.
        :rtype: str
        """
        text = DocumentParser._parse_pdf_pdfplumber(file_path)
        if text:
            return text

        logger.warning(
            "pdfplumber returned empty text for '%s'; trying PyPDF2.", file_path
        )
        text = DocumentParser._parse_pdf_pypdf2(file_path)
        if text:
            return text

        logger.error(
            "Both pdfplumber and PyPDF2 failed to extract text from '%s'.",
            file_path,
        )
        return ""

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Iterates paragraphs and table cells (skills sections are frequently
        formatted as tables in professional CVs).

        :param file_path: Path to the DOCX file.
        :type file_path: str
        :raises ImportError: python-docx is not installed.
        :raises ValueError: File cannot be parsed.
        """
        try:
            from docx import Document  # noqa: PLC0415
        except ImportError as exc:
            raise ImportError(
                "python-docx is not installed. Run: pip install python-docx"
            ) from exc

        try:
            doc = Document(file_path)
            parts: list[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())

            return "\n".join(parts)
        except Exception as exc:
            logger.exception("Failed to parse DOCX '%s'.", file_path)
            raise ValueError(f"Could not parse DOCX file: {exc}") from exc

                                                                        
             
                                                                        

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Return True if the file extension is handled by this parser."""
        return Path(filename).suffix.lower() in {".pdf", ".docx"}

    @staticmethod
    def clean_text(raw: str) -> str:
        """
        Normalise whitespace in extracted text before NLP processing.

        - Strips control characters (keeps tabs/spaces).
        - Collapses consecutive blank lines to a single blank line.
        - Strips leading/trailing whitespace per line.

        :param raw: Raw extracted text.
        :type raw: str
        :returns: Cleaned text.
        :rtype: str
        """
        import unicodedata                             

        lines: list[str] = []
        prev_blank = False

        for line in raw.splitlines():
            cleaned = "".join(
                ch
                for ch in line
                if unicodedata.category(ch) not in {"Cc", "Cf"} or ch in "\t "
            ).strip()

            if cleaned:
                lines.append(cleaned)
                prev_blank = False
            else:
                if not prev_blank:
                    lines.append("")
                prev_blank = True

        return "\n".join(lines).strip()

                                                                        
                     
                                                                        

    @staticmethod
    def _parse_pdf_pdfplumber(file_path: str) -> str:
        """Primary PDF extractor using pdfplumber."""
        try:
            import pdfplumber  # noqa: PLC0415
        except ImportError:
            logger.warning("pdfplumber not installed; skipping primary PDF extractor.")
            return ""

        try:
            pages: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
            return "\n".join(pages).strip()
        except Exception as exc:
            logger.warning("pdfplumber error on '%s': %s", file_path, exc)
            return ""

    @staticmethod
    def _parse_pdf_pypdf2(file_path: str) -> str:
        """Fallback PDF extractor using PyPDF2."""
        try:
            import PyPDF2  # noqa: PLC0415
        except ImportError:
            logger.warning("PyPDF2 not installed; fallback PDF extractor unavailable.")
            return ""

        try:
            pages: list[str] = []
            with open(file_path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages).strip()
        except Exception as exc:
            logger.warning("PyPDF2 error on '%s': %s", file_path, exc)
            return ""

    @staticmethod
    def _parse_doc_unsupported(_file_path: str) -> str:
        raise NotImplementedError(
            "Legacy .doc format is not supported. "
            "Please convert the file to .docx or .pdf before uploading, "
            "or install the 'textract' library and extend this method."
        )